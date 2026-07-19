"""
DOCX text extraction module.
"""

import subprocess
import sys
from typing import Any, Dict, List
import docx

def extract_doc_text(file_path: str) -> List[Dict[str, Any]]:
    """Extract text from an older binary .doc file.
    
    On macOS, uses the built-in `textutil` command.
    On other platforms, attempts to use `antiword` if installed.
    """
    if sys.platform == "darwin":
        try:
            result = subprocess.run(
                ["textutil", "-convert", "txt", "-stdout", file_path],
                capture_output=True,
                text=True,
                check=True
            )
            return [{"page": 1, "text": result.stdout}]
        except subprocess.CalledProcessError as exc:
            raise ValueError(f"textutil failed to read .doc file: {exc.stderr}")
        except FileNotFoundError:
            pass # Fall through if textutil is missing for some reason

    # Fallback for Linux/Windows: try antiword
    try:
        result = subprocess.run(
            ["antiword", file_path],
            capture_output=True,
            text=True,
            check=True
        )
        return [{"page": 1, "text": result.stdout}]
    except (subprocess.CalledProcessError, FileNotFoundError) as exc:
        raise ValueError(f"Failed to read .doc file {file_path}. On macOS, textutil is required. On Linux/Windows, antiword is required. Error: {exc}")


def extract_docx_text(file_path: str) -> List[Dict[str, Any]]:
    """Extract text from a DOCX file.
    
    Returns a list of dictionaries with page=1 and text content.
    Currently, python-docx does not support pagination natively, so we return
    all text as a single page.
    """
    try:
        doc = docx.Document(file_path)
    except Exception as exc:
        raise ValueError(f"Failed to read docx file {file_path}: {exc}")
        
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
        
    # We can also extract text from tables if needed
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                row_text.append(cell.text)
            full_text.append(" | ".join(row_text))
            
    text_content = "\n".join(full_text)
    
    return [{"page": 1, "text": text_content}]
